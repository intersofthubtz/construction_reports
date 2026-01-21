import datetime
from io import BytesIO
import csv
from itertools import groupby
import os
from django.conf import settings
from django.contrib import messages
from openpyxl.utils import get_column_letter
from django.contrib.auth.decorators import login_required, permission_required
from django.http import HttpResponse
from django.shortcuts import render, redirect, get_object_or_404
from compliance.models import Compliance
from projects.models import Project
from finance.models import PaymentCertificate, FundTransaction
from resources.models import Equipment, Manpower
from quality.models import MaterialTest, WorkApproval
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment
from reportlab.platypus import SimpleDocTemplate, PageBreak, Table, TableStyle, Image, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
import logging
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns
from sitemanage.models import Activity, ProgressLog
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reports.models import ProgressReportCover
from reports.forms import ProgressReportCoverForm
from django.core.paginator import Paginator


logger = logging.getLogger(__name__)

# ---------------- Helpers ----------------
def get_allowed_projects(user):
    """Return active projects the user can access."""
    if user.is_superuser or user.is_staff:
        return Project.objects.filter(is_active=True)
    return Project.objects.filter(
        is_active=True,
        participants__user=user,
        participants__is_active=True
    ).distinct()

def filter_by_allowed_projects(queryset, user, project_field="project"):
    """Filter queryset by allowed projects"""
    if user.is_superuser or user.is_staff:
        return queryset
    return queryset.filter(**{f"{project_field}__in": get_allowed_projects(user)})


def _validate_progress_download(request):
    project_id = request.GET.get("project")
    if not project_id:
        messages.error(request, "Please select a project before downloading the report.")
        return False
    return True


def safe(value):
    return value.strftime("%Y-%m-%d") if value else "-"


# ---------------- Views ----------------
@login_required
@permission_required("reports.view_progressreportcover", raise_exception=True)
def progress_cover_list(request):
    search = request.GET.get("q", "")
    covers = ProgressReportCover.objects.select_related("project").order_by("-created_at")
    covers = filter_by_allowed_projects(covers, request.user)

    if search:
        covers = covers.filter(
            project__project_name__icontains=search
        ) | covers.filter(report_no__icontains=search)

    paginator = Paginator(covers, 10)
    page_obj = paginator.get_page(request.GET.get("page"))

    return render(request, "reports/progress_cover/list.html", {
        "covers": page_obj,
        "page_obj": page_obj,
        "search": search,
    })


@login_required
@permission_required("reports.add_progressreportcover", raise_exception=True)
def progress_cover_create(request):
    """Create a new progress report cover (only allowed projects)."""
    form = ProgressReportCoverForm(request.POST or None, request.FILES or None, user=request.user)

    if request.method == "POST" and form.is_valid():
        cover = form.save(commit=False)
        cover.created_by = request.user
        cover.save()
        messages.success(request, "Progress report cover created successfully!")
        return redirect("reports:progress_cover_list")

    return render(request, "reports/progress_cover/form.html", {
        "form": form,
        "title": "Generate Progress Report Cover",
        "is_edit": False,
    })


@login_required
@permission_required("reports.change_progressreportcover", raise_exception=True)
def progress_cover_edit(request, pk):
    """Edit an existing cover (only allowed projects)."""
    cover = get_object_or_404(
        ProgressReportCover,
        pk=pk,
        project__in=get_allowed_projects(request.user)
    )

    form = ProgressReportCoverForm(request.POST or None, request.FILES or None, instance=cover, user=request.user)

    if request.method == "POST" and form.is_valid():
        # Preserve old image if not replaced
        if 'cover_image' in request.FILES:
            cover.cover_image = request.FILES['cover_image']
        form.save()
        messages.success(request, "Progress report cover updated successfully!")
        return redirect("reports:progress_cover_list")

    return render(request, "reports/progress_cover/form.html", {
        "form": form,
        "title": "Edit Progress Report Cover",
        "is_edit": True,
    })


@login_required
@permission_required("reports.delete_progressreportcover", raise_exception=True)
def progress_cover_delete(request, pk):
    """Delete a cover (only allowed projects)."""
    cover = get_object_or_404(
        ProgressReportCover,
        pk=pk,
        project__in=get_allowed_projects(request.user)
    )

    if request.method == "POST":
        cover.delete()
        messages.success(request, "Progress report cover deleted successfully!")
        return redirect("reports:progress_cover_list")

    return render(request, "reports/progress_cover/confirm_delete.html", {"cover": cover})


@login_required
@permission_required("reports.view_progressreportcover", raise_exception=True)
def progress_cover_pdf(request, pk):
    """
    Generate PDF for a Progress Report Cover.
    """
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    import os

    # Fetch cover, restrict to allowed projects
    cover = get_object_or_404(
        ProgressReportCover,
        pk=pk,
        project__in=get_allowed_projects(request.user)
    )

    # --- HTTP Response ---
    response = HttpResponse(content_type="application/pdf")
    filename = f"Progress_Report_{cover.report_no}.pdf"
    response["Content-Disposition"] = f'inline; filename="{filename}"'

    # --- PDF Document Setup ---
    doc = SimpleDocTemplate(
        response,
        pagesize=A4,
        rightMargin=2*cm, leftMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm
    )

    elements = []
    styles = getSampleStyleSheet()

    # --- Styles ---
    project_style = ParagraphStyle(
        "project", parent=styles["Title"], alignment=1, fontSize=22, spaceAfter=6
    )
    title_style = ParagraphStyle(
        "title", parent=styles["Title"], alignment=1, fontSize=18, spaceAfter=12
    )
    normal_centered = ParagraphStyle(
        "normal_centered", parent=styles["Normal"], alignment=1, fontSize=14, spaceAfter=6
    )

    # --- Logo ---
    logo_path = os.path.join(settings.BASE_DIR, "static/images/nhc_logo.jpg")
    if os.path.exists(logo_path):
        logo = Image(logo_path, width=5*cm, height=5*cm)
        logo.hAlign = 'CENTER'
        elements.append(logo)
        elements.append(Spacer(1, 1*cm))
    else:
        elements.append(Paragraph("NHC LOGO MISSING", project_style))
        elements.append(Spacer(1, 1*cm))

    # --- Project Name & Report Title ---
    elements.append(Paragraph(f"<b>{cover.project.project_name.upper()}</b>", project_style))
    elements.append(Paragraph(f"<b>{cover.report_title.upper()}</b>", title_style))

    # --- Cover Image (Optional) ---
    if cover.cover_image and os.path.exists(cover.cover_image.path):
        img = Image(cover.cover_image.path, width=14*cm, height=8*cm)
        img.hAlign = 'CENTER'
        elements.append(img)
        elements.append(Spacer(1, 0.8*cm))

    # --- Report Period ---
    period_text = f"FROM {cover.period_from.strftime('%d %B %Y')} TO {cover.period_to.strftime('%d %B %Y')}"
    elements.append(Paragraph(period_text, normal_centered))

    # --- Prepared By ---
    prepared_by_text = f"<b>PREPARED BY:</b> {cover.prepared_by}"
    elements.append(Paragraph(prepared_by_text, normal_centered))

    # --- Build PDF ---
    doc.build(elements)

    return response



# ---------------- Project Report ----------------
@login_required
@permission_required("reports.view_projectreport", raise_exception=True)
def project_report(request):
    try:
        projects = get_allowed_projects(request.user)
        project_id = request.GET.get("project")
        from_date = request.GET.get("from_date")
        to_date = request.GET.get("to_date")

        is_filtered = bool(project_id or from_date or to_date)

        queryset = Project.objects.none()

        if is_filtered:
            queryset = filter_by_allowed_projects(
                Project.objects.filter(is_active=True),
                request.user,
                project_field="id"
            )

            if project_id and project_id not in ["", "None"]:
                queryset = queryset.filter(id=project_id)
            if from_date:
                queryset = queryset.filter(commencement_date__gte=from_date)
            if to_date:
                queryset = queryset.filter(practical_completion_date__lte=to_date)

        context = {
            "projects": projects,
            "project_list": queryset.order_by("project_name"),
            "filter_project": project_id,
            "filter_from": from_date,
            "filter_to": to_date,
            "is_filtered": is_filtered,
        }
        return render(request, "reports/project/list.html", context)

    except Exception as e:
        logger.exception("Project report error")
        messages.error(request, "Unable to load Project Report.")
        return redirect("reports:project_report")

# ----------------------------
# Excel Download
# ----------------------------
@login_required
@permission_required("reports.view_projectreport", raise_exception=True)
def project_report_download_excel(request):
    project_id = request.GET.get("project")

    # Validation
    if not project_id or project_id in ["None", ""]:
        messages.error(request, "Please select a project before downloading.")
        return redirect("reports:project_report")

    projects = filter_by_allowed_projects(
        Project.objects.filter(is_active=True, id=project_id), request.user, project_field="id"
    )

    wb = Workbook()
    ws = wb.active
    ws.title = "Project Report"

    headers = [
        "Project Code", "Project Name", "Client", "Location",
        "Contract Sum", "Contract Duration (Months)", "Contract Signing Date",
        "Mobilization Start", "Mobilization End", "Commencement Date", "Practical Completion",
        "Delay Status", "Defects Liability (Days)", "Defects Start", "Defects End",
        "Participants", "Contractors"
    ]

    ws.append(headers)

    for p in projects:
        participants = ", ".join([
            f"{part.user.get_full_name() or part.user.username} ({part.project_role.name})"
            for part in p.participants.filter(is_active=True)
        ]) or "N/A"

        contractors = ", ".join([
            f"{c.contractor.name} ({c.contractor.contractor_type.name}) - {c.work_description}"
            for c in p.contractors.filter(is_active=True)
        ]) or "N/A"

        ws.append([
            p.project_code, p.project_name, p.client.name, p.location,
            float(p.contract_sum), p.contract_duration_months,
            p.contract_signing_date.strftime("%Y-%m-%d"),
            p.mobilization_start.strftime("%Y-%m-%d"),
            p.mobilization_end.strftime("%Y-%m-%d"),
            p.commencement_date.strftime("%Y-%m-%d"),
            p.practical_completion_date.strftime("%Y-%m-%d"),
            p.delay_status, p.defects_liability_period_days,
            p.defects_start.strftime("%Y-%m-%d") if p.defects_start else "",
            p.defects_end.strftime("%Y-%m-%d") if p.defects_end else "",
            participants, contractors
        ])

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    response = HttpResponse(
        buffer, content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    response["Content-Disposition"] = f'attachment; filename="project_report_{datetime.date.today()}.xlsx"'
    return response


# ----------------------------
# PDF Download
# ----------------------------

@login_required
@permission_required("reports.view_projectreport", raise_exception=True)
def project_report_download_pdf(request):
    project_id = request.GET.get("project")

    if not project_id:
        messages.error(request, "Please select a project before downloading.")
        return redirect("reports:project_report")

    projects = filter_by_allowed_projects(
        Project.objects.filter(is_active=True, id=project_id),
        request.user,
        project_field="id"
    )

    buffer = BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=36
    )

    styles = getSampleStyleSheet()
    elements = []

    # ---------- Styles ----------
    title_style = ParagraphStyle(
        "Title",
        parent=styles["Heading1"],
        alignment=1,
        spaceAfter=20
    )

    section_style = ParagraphStyle(
        "Section",
        parent=styles["Heading2"],
        textColor=colors.HexColor("#0F5391"),
        spaceBefore=14,
        spaceAfter=8
    )

    cell_style = ParagraphStyle(
        "Cell",
        parent=styles["Normal"],
        fontSize=9,
        leading=12
    )

    # ---------- Title ----------
    elements.append(Paragraph("PROJECT OVERVIEW", title_style))

    for p in projects:

        elements.append(Paragraph(
            f"{p.project_name}",
            section_style
        ))

        # =======================
        # PROJECT DETAILS (ONE TABLE)
        # =======================

        project_details = [
            ("Project Code", p.project_code),
            ("Project Name", p.project_name),
            ("Client", p.client.name),
            ("Location", p.location),
            ("Contract Sum", f"{p.contract_sum:,.2f}"),
            ("Contract Duration (Months)", p.contract_duration_months),
            ("Contract Signing Date", safe(p.contract_signing_date)),
            ("Site Possession Date", safe(p.site_possession_date)),
            ("Mobilization Start", safe(p.mobilization_start)),
            ("Mobilization End", safe(p.mobilization_end)),
            ("Commencement Date", safe(p.commencement_date)),
            ("Practical Completion Date", safe(p.practical_completion_date)),
            ("Delay Status", p.delay_status),
            ("Defects Liability Period (Days)", p.defects_liability_period_days),
            ("Defects Start Date", safe(p.defects_start)),
            ("Defects End Date", safe(p.defects_end)),
        ]

        project_table = Table(
            [[
                Paragraph(label, cell_style),
                Paragraph(str(value), cell_style)
            ] for label, value in project_details],
            colWidths=[200, doc.width - 200]
        )

        project_table.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))

        elements.append(project_table)
        elements.append(Spacer(1, 16))

        # =======================
        # PARTICIPANTS
        # =======================

        elements.append(Paragraph("Project Participants", section_style))

        participant_data = [["Name", "Role"]]

        for part in p.participants.filter(is_active=True):
            participant_data.append([
                Paragraph(part.user.get_full_name() or part.user.username, cell_style),
                Paragraph(part.project_role.name, cell_style),
            ])

        if len(participant_data) == 1:
            participant_data.append([
                Paragraph("No participants assigned", cell_style),
                Paragraph("-", cell_style),
            ])

        participant_table = Table(
            participant_data,
            colWidths=[doc.width * 0.6, doc.width * 0.4],
            repeatRows=1
        )

        participant_table.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))

        elements.append(participant_table)
        elements.append(Spacer(1, 16))

        # =======================
        # CONTRACTORS (WRAP SAFE)
        # =======================

        elements.append(Paragraph("Contractors", section_style))

        contractor_data = [["Contractor", "Type", "Work Description"]]

        for c in p.contractors.filter(is_active=True):
            contractor_data.append([
                Paragraph(c.contractor.name, cell_style),
                Paragraph(c.contractor.contractor_type.name, cell_style),
                Paragraph(c.work_description, cell_style),  # WRAPS SAFELY
            ])

        if len(contractor_data) == 1:
            contractor_data.append([
                Paragraph("No contractors assigned", cell_style),
                Paragraph("-", cell_style),
                Paragraph("-", cell_style),
            ])

        contractor_table = Table(
            contractor_data,
            colWidths=[
                doc.width * 0.25,
                doc.width * 0.20,
                doc.width * 0.55,
            ],
            repeatRows=1
        )

        contractor_table.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))

        elements.append(contractor_table)
        elements.append(Spacer(1, 28))

    # ---------- Build ----------
    doc.build(elements)
    buffer.seek(0)

    response = HttpResponse(buffer, content_type="application/pdf")
    response["Content-Disposition"] = (
        f'attachment; filename="project_report_{datetime.date.today()}.pdf"'
    )
    return response




# ------------------ HELPERS ------------------

def remove_paragraph_spacing(paragraph):
    """Remove spacing before and after paragraph."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(ns.qn("w:before"), "0")
    spacing.set(ns.qn("w:after"), "0")
    spacing.set(ns.qn("w:line"), "240")
    spacing.set(ns.qn("w:lineRule"), "auto")
    pPr.append(spacing)

def normalize_cell(cell, bold=False):
    """Set cell text formatting and remove spacing."""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if p.runs:
        p.runs[0].bold = bold
    remove_paragraph_spacing(p)

def style_header_cell(cell):
    """Style table header cell (center + bold)."""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(10)

def set_fixed_table_layout(table):
    """Fix column widths for Table Grid."""
    for row in table.rows:
        for cell in row.cells:
            cell.width = cell.width

def safe_text(value):
    """Convert None values to '-' and format dates."""
    if value is None:
        return "-"
    if hasattr(value, "strftime"):
        return value.strftime("%Y-%m-%d")
    return str(value)

@login_required
@permission_required("reports.view_projectreport", raise_exception=True)
def project_report_download_word(request):
    project_id = request.GET.get("project")
    if not project_id:
        messages.error(request, "Please select a project before downloading.")
        return redirect("reports:project_report")

    projects = filter_by_allowed_projects(
        Project.objects.filter(is_active=True, id=project_id),
        request.user,
        project_field="id"
    )

    doc = Document()

    # ------------------ PAGE SETUP ------------------
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    LABEL_COL = Inches(2.5)
    VALUE_COL = Inches(4.77)

    # ------------------ TITLE ------------------
    title = doc.add_heading("PROJECT REPORT", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    remove_paragraph_spacing(title)

    for p in projects:
        # ------------------ PROJECT HEADING ------------------
        heading = doc.add_heading(p.project_name, level=2)
        heading.runs[0].bold = True
        remove_paragraph_spacing(heading)

        # ------------------ PROJECT DETAILS ------------------
        details = [
            ("Project Code", p.project_code),
            ("Client", p.client.name),
            ("Location", p.location),
            ("Contract Sum", f"{p.contract_sum:,.2f}"),
            ("Contract Duration (Months)", p.contract_duration_months),
            ("Commencement Date", p.commencement_date),
            ("Practical Completion Date", p.practical_completion_date),
            ("Delay Status", p.delay_status),
            ("Defects Liability Period (Days)", p.defects_liability_period_days),
        ]

        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = LABEL_COL
        table.columns[1].width = VALUE_COL

        for label, value in details:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = safe_text(value)
            normalize_cell(row[0], bold=True)
            normalize_cell(row[1])

        remove_paragraph_spacing(doc.paragraphs[-1])

        # ------------------ PROJECT PARTICIPANTS ------------------
        doc.add_heading("Project Participants", level=3).runs[0].bold = True
        remove_paragraph_spacing(doc.paragraphs[-1])

        pt = doc.add_table(rows=1, cols=2)
        pt.style = "Table Grid"
        pt.alignment = WD_TABLE_ALIGNMENT.CENTER
        pt.autofit = False
        pt.columns[0].width = LABEL_COL
        pt.columns[1].width = VALUE_COL

        hdr = pt.rows[0].cells
        hdr[0].text = "Participant"
        hdr[1].text = "Role"
        normalize_cell(hdr[0], bold=True)
        normalize_cell(hdr[1], bold=True)

        participants = p.participants.filter(is_active=True)
        if participants.exists():
            for part in participants:
                row = pt.add_row().cells
                row[0].text = part.user.get_full_name() or part.user.username
                row[1].text = part.project_role.name
                normalize_cell(row[0])
                normalize_cell(row[1])
        else:
            row = pt.add_row().cells
            row[0].text = "No participants assigned"
            row[1].text = "-"
            normalize_cell(row[0])
            normalize_cell(row[1])

        remove_paragraph_spacing(doc.paragraphs[-1])

        # ------------------ CONTRACTORS ------------------
        doc.add_heading("Contractors", level=3).runs[0].bold = True
        remove_paragraph_spacing(doc.paragraphs[-1])

        ct = doc.add_table(rows=1, cols=2)
        ct.style = "Table Grid"
        ct.alignment = WD_TABLE_ALIGNMENT.CENTER
        ct.autofit = False
        ct.columns[0].width = LABEL_COL
        ct.columns[1].width = VALUE_COL

        hdr = ct.rows[0].cells
        hdr[0].text = "Contractor"
        hdr[1].text = "Details"
        normalize_cell(hdr[0], bold=True)
        normalize_cell(hdr[1], bold=True)

        contractors = p.contractors.filter(is_active=True)
        if contractors.exists():
            for c in contractors:
                row = ct.add_row().cells
                row[0].text = c.contractor.name
                row[1].text = (
                    f"Type: {c.contractor.contractor_type.name} | "
                    f"Work: {c.work_description}"
                )
                normalize_cell(row[0])
                normalize_cell(row[1])
        else:
            row = ct.add_row().cells
            row[0].text = "No contractors assigned"
            row[1].text = "-"
            normalize_cell(row[0])
            normalize_cell(row[1])

        remove_paragraph_spacing(doc.paragraphs[-1])

    # ------------------ EXPORT ------------------
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer,
        content_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        )
    )
    response["Content-Disposition"] = (
        f'attachment; filename="project_report_{datetime.date.today()}.docx"'
    )
    return response



# ------------------ PROGRESS REPORT VIEW ------------------
@login_required
@permission_required("reports.view_progressreport", raise_exception=True)
def progress_report(request):
    try:
        projects = get_allowed_projects(request.user)
        project_id = request.GET.get("project")
        from_date = request.GET.get("from_date")
        to_date = request.GET.get("to_date")

        is_filtered = bool(project_id or from_date or to_date)

        activities = Activity.objects.none()
        progress_logs = ProgressLog.objects.none()

        if is_filtered:
            activities = filter_by_allowed_projects(
                Activity.objects.filter(is_active=True),
                request.user,
                project_field="project"
            )

            progress_logs = filter_by_allowed_projects(
                ProgressLog.objects.filter(is_active=True),
                request.user,
                project_field="activity__project"
            )

            if project_id:
                activities = activities.filter(project_id=project_id)
                progress_logs = progress_logs.filter(activity__project_id=project_id)
            if from_date:
                progress_logs = progress_logs.filter(date__gte=from_date)
            if to_date:
                progress_logs = progress_logs.filter(date__lte=to_date)

        context = {
            "projects": projects,
            "activities": activities,
            "progress_logs": progress_logs,
            "filter_project": project_id,
            "filter_from": from_date,
            "filter_to": to_date,
            "is_filtered": is_filtered,
        }
        return render(request, "reports/progress/list.html", context)

    except Exception as e:
        logger.exception("Progress report error")
        messages.error(request, "Unable to load Progress Report.")
        return redirect("reports:progress_report")


@login_required
@permission_required("reports.view_progressreport", raise_exception=True)
def progress_report_download_excel(request):

    if not _validate_progress_download(request):
        return redirect("reports:progress_report")

    activities = filter_by_allowed_projects(
        Activity.objects.filter(is_active=True),
        request.user,
        "project"
    )

    logs = filter_by_allowed_projects(
        ProgressLog.objects.filter(is_active=True),
        request.user,
        "activity__project"
    )

    project_id = request.GET.get("project")
    from_date = request.GET.get("from_date")
    to_date = request.GET.get("to_date")

    activities = activities.filter(project_id=project_id)
    logs = logs.filter(activity__project_id=project_id)

    if from_date:
        logs = logs.filter(date__gte=from_date)
    if to_date:
        logs = logs.filter(date__lte=to_date)

    response = HttpResponse(content_type="text/csv")
    response["Content-Disposition"] = 'attachment; filename="progress_report.csv"'
    writer = csv.writer(response)

    writer.writerow(["ACTIVITIES"])
    writer.writerow(["Project", "Activity", "Status", "Progress %", "Planned End"])

    for a in activities:
        writer.writerow([
            a.project.project_name,
            a.name,
            a.status,
            a.progress_percent,
            a.planned_end
        ])

    writer.writerow([])
    writer.writerow(["PROGRESS LOGS"])
    writer.writerow(["Activity", "Date", "Progress %", "Remarks"])

    for log in logs:
        writer.writerow([
            log.activity.name,
            log.date,
            log.progress_percent,
            log.remarks
        ])

    return response



@login_required
@permission_required("reports.view_progressreport", raise_exception=True)
def progress_report_download_pdf(request):
    project_id = request.GET.get("project")
    if not project_id:
        from django.contrib import messages
        messages.error(request, "Please select a project before downloading.")
        return redirect("reports:progress_report")

    # Filter logs and activities
    logs = filter_by_allowed_projects(
        ProgressLog.objects.filter(is_active=True),
        request.user,
        project_field="activity__project"
    ).filter(activity__project_id=project_id).select_related(
        "activity", "activity__category"
    ).order_by("activity__category__name", "activity__name", "date")

    if not logs.exists():
        from django.contrib import messages
        messages.error(request, "No activities found for the selected project.")
        return redirect("reports:progress_report")

    project = logs.first().activity.project

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=36
    )
    styles = getSampleStyleSheet()
    elements = []

    # ---------- Styles ----------
    title_style = ParagraphStyle(
        "Title", parent=styles["Heading1"], alignment=1, spaceAfter=10
    )
    section_style = ParagraphStyle(
        "Section", parent=styles["Heading2"], textColor=colors.HexColor("#0F5391"),
        spaceBefore=10, spaceAfter=5
    )
    cell_style = ParagraphStyle(
        "Cell", parent=styles["Normal"], fontSize=9, leading=12
    )
    img_caption_style = ParagraphStyle(
        "ImgCaption", parent=styles["Normal"], fontSize=8, leading=10, alignment=1
    )

    # ---------- Title ----------
    elements.append(Paragraph(f"{project.project_name} <font size=9>({project.project_code})</font>", title_style))

    # -------------------------
    # Activities grouped by category
    # -------------------------
    activities = Activity.objects.filter(project=project, is_active=True).select_related("category").order_by("category__name", "name")
    grouped_activities = groupby(activities, lambda a: a.category.name if a.category else "Uncategorized")

    all_activity_images = []

    for category_name, acts in grouped_activities:
        acts_list = list(acts)
        elements.append(Paragraph(f"Category: {category_name}", section_style))

        # Activity Table per category
        table_data = [["S/N", "Activity Description", "Progress %", "Remarks"]]
        for idx, act in enumerate(acts_list, start=1):
            log = logs.filter(activity=act).last()
            remark = log.remarks if log else "-"
            table_data.append([
                Paragraph(str(idx), cell_style),
                Paragraph(act.name, cell_style),
                Paragraph(f"{act.progress_percent}%", cell_style),
                Paragraph(remark, cell_style)
            ])
            all_activity_images.extend(list(act.activity_images.filter(is_active=True)))

        col_widths = [40, doc.width - 180, 60, 80]
        act_table = Table(table_data, colWidths=col_widths, repeatRows=1)
        act_table.setStyle(TableStyle([
            ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
            ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("ALIGN", (0,0), (0,-1), "CENTER"),
            ("ALIGN", (2,1), (2,-1), "CENTER"),
            ("ALIGN", (3,1), (3,-1), "CENTER"),
        ]))
        elements.append(act_table)

    # -------------------------
    # STATUS OF ON-GOING SITE WORKS (Only In Progress)
    # -------------------------
    elements.append(Spacer(1, 5))
    elements.append(Paragraph("STATUS OF ON-GOING SITE WORKS", section_style))

    ongoing_activities = [act for act in activities if act.status == "In Progress"]
    if ongoing_activities:
        status_table_data = [["S/N", "Activity", "Status", "Progress %"]]
        for idx, act in enumerate(ongoing_activities, start=1):
            status_table_data.append([
                Paragraph(str(idx), cell_style),
                Paragraph(act.name, cell_style),
                Paragraph(act.status, cell_style),
                Paragraph(f"{act.progress_percent}%", cell_style)
            ])
        col_widths = [40, doc.width - 180, 80, 60]
        status_table = Table(status_table_data, colWidths=col_widths, repeatRows=1)
        status_table.setStyle(TableStyle([
            ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
            ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("ALIGN", (0,0), (0,-1), "CENTER"),
            ("ALIGN", (3,1), (3,-1), "CENTER"),
        ]))
        elements.append(status_table)
    else:
        elements.append(Paragraph("No activities currently in progress.", cell_style))

    # -------------------------
    # ALL ACTIVITY IMAGES THUMBNAILS (After Tables)
    # -------------------------
    if all_activity_images:
        elements.append(PageBreak())
        elements.append(Paragraph("SITE IMAGES FOR ACTIVITIES", section_style))
        img_width = (doc.width - 40) / 3
        img_height = img_width * 0.75
        row_imgs = []

        for idx, img_obj in enumerate(all_activity_images, start=1):
            try:
                im = Image(img_obj.image.path, width=img_width, height=img_height, kind='proportional')
                caption = Paragraph(f"{img_obj.activity.name} ({safe(img_obj.image_date)})", img_caption_style)
                row_imgs.append([im, Spacer(1,2), caption])
            except Exception:
                continue

            if idx % 3 == 0:
                t = Table([row_imgs], colWidths=[img_width]*len(row_imgs))
                t.setStyle(TableStyle([
                    ("VALIGN", (0,0), (-1,-1), "TOP"),
                    ("ALIGN", (0,0), (-1,-1), "CENTER"),
                ]))
                elements.append(t)
                row_imgs = []

        if row_imgs:
            t = Table([row_imgs], colWidths=[img_width]*len(row_imgs))
            t.setStyle(TableStyle([
                ("VALIGN", (0,0), (-1,-1), "TOP"),
                ("ALIGN", (0,0), (-1,-1), "CENTER"),
            ]))
            elements.append(t)

    # -------------------------
    # Build PDF
    # -------------------------
    doc.build(elements)
    buffer.seek(0)
    response = HttpResponse(buffer, content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="progress_report_{datetime.date.today()}.pdf"'
    return response


# ------------------ HELPERS ------------------

def remove_paragraph_spacing(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()

    spacing = OxmlElement('w:spacing')
    spacing.set(ns.qn('w:before'), '0')
    spacing.set(ns.qn('w:after'), '0')
    spacing.set(ns.qn('w:line'), '240')
    spacing.set(ns.qn('w:lineRule'), 'auto')

    pPr.append(spacing)


def style_header_cell(cell):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(10)


# ------------------ VIEW ------------------

@login_required
@permission_required("reports.view_progressreport", raise_exception=True)
def progress_report_download_word(request):

    if not _validate_progress_download(request):
        return redirect("reports:progress_report")

    project_id = request.GET.get("project")

    logs = filter_by_allowed_projects(
        ProgressLog.objects.filter(is_active=True),
        request.user,
        "activity__project"
    ).filter(
        activity__project_id=project_id
    ).select_related(
        "activity", "activity__category"
    ).order_by(
        "activity__category__name", "activity__name", "date"
    )

    if not logs.exists():
        messages.error(request, "No activities found for the selected project.")
        return redirect("reports:progress_report")

    project = logs.first().activity.project

    doc = Document()

    # ------------------ TITLE ------------------

    title = doc.add_heading("PROJECT PROGRESS REPORT", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    remove_paragraph_spacing(title)

    subtitle = doc.add_paragraph(f"{project.project_name} ({project.project_code})")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    remove_paragraph_spacing(subtitle)

    # ------------------ ACTIVITIES BY CATEGORY ------------------

    activities = Activity.objects.filter(
        project=project,
        is_active=True
    ).select_related("category").order_by("category__name", "name")

    for category_name, acts in groupby(
        activities,
        lambda a: a.category.name if a.category else "Uncategorized"
    ):
        # Category heading
        heading = doc.add_heading(f"Category: {category_name}", level=2)
        heading.runs[0].bold = True
        remove_paragraph_spacing(heading)

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ["S/N", "Activity Description", "Progress %", "Remarks"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            style_header_cell(table.rows[0].cells[i])

        for idx, act in enumerate(list(acts), start=1):
            log = logs.filter(activity=act).last()
            row = table.add_row().cells

            row[0].text = str(idx)
            row[1].text = act.name
            row[2].text = f"{act.progress_percent}%"
            row[3].text = log.remarks if log and log.remarks else "-"

        # Remove spacing after table
        remove_paragraph_spacing(doc.paragraphs[-1])

    # ------------------ STATUS OF ON-GOING SITE WORKS ------------------

    status_heading = doc.add_heading("STATUS OF ON-GOING SITE WORKS", level=2)
    status_heading.runs[0].bold = True
    remove_paragraph_spacing(status_heading)

    ongoing_activities = activities.filter(status="In Progress")

    if ongoing_activities.exists():
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ["S/N", "Activity Name", "Status", "Progress %"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            style_header_cell(table.rows[0].cells[i])

        for idx, act in enumerate(ongoing_activities, start=1):
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = act.name
            row[2].text = act.status
            row[3].text = f"{act.progress_percent}%"

        remove_paragraph_spacing(doc.paragraphs[-1])

    else:
        p = doc.add_paragraph("No activities currently in progress.")
        remove_paragraph_spacing(p)

    # ------------------ EXPORT ------------------

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = (
        f'attachment; filename="progress_report_{project.project_code}.docx"'
    )

    return response





# ------------------ RESOURCES REPORT VIEW ------------------
@login_required
@permission_required("reports.view_resourcesreport", raise_exception=True)
def resources_report(request):
    try:
        projects = get_allowed_projects(request.user)
        project_id = request.GET.get("project")
        from_date = request.GET.get("from_date")
        to_date = request.GET.get("to_date")

        is_filtered = bool(project_id or from_date or to_date)

        equipment = Equipment.objects.none()
        manpower = Manpower.objects.none()

        if is_filtered:
            equipment = filter_by_allowed_projects(
                Equipment.objects.filter(is_active=True),
                request.user,
                project_field="project"
            )
            manpower = filter_by_allowed_projects(
                Manpower.objects.filter(is_active=True),
                request.user,
                project_field="project"
            )

            if project_id:
                equipment = equipment.filter(project_id=project_id)
                manpower = manpower.filter(project_id=project_id)
            if from_date:
                equipment = equipment.filter(delivery_date__gte=from_date)
                manpower = manpower.filter(start_date__gte=from_date)
            if to_date:
                equipment = equipment.filter(delivery_date__lte=to_date)
                manpower = manpower.filter(start_date__lte=to_date)

        context = {
            "projects": projects,
            "equipment": equipment,
            "manpower": manpower,
            "filter_project": project_id,
            "filter_from": from_date,
            "filter_to": to_date,
            "is_filtered": is_filtered,
        }
        return render(request, "reports/resources/list.html", context)

    except Exception:
        logger.exception("Resources report error")
        messages.error(request, "Unable to load Resources Report.")
        return redirect("reports:resources_report")


# ------------------ EXCEL ------------------
@login_required
@permission_required("reports.view_resourcesreport", raise_exception=True)
def resources_report_download_excel(request):
    project_id = request.GET.get("project")
    from_date = request.GET.get("from_date")
    to_date = request.GET.get("to_date")

    if not project_id or project_id in ["None", ""]:
        messages.error(request, "Please select a project before downloading.")
        return redirect("reports:resources_report")

    try:
        equipment = filter_by_allowed_projects(
            Equipment.objects.filter(is_active=True).order_by("project", "name"), request.user, project_field="project"
        )
        manpower = filter_by_allowed_projects(
            Manpower.objects.filter(is_active=True).order_by("project", "role"), request.user, project_field="project"
        )

        if project_id:
            equipment = equipment.filter(project_id=int(project_id))
            manpower = manpower.filter(project_id=int(project_id))
        if from_date:
            equipment = equipment.filter(delivery_date__gte=from_date)
            manpower = manpower.filter(start_date__gte=from_date)
        if to_date:
            equipment = equipment.filter(delivery_date__lte=to_date)
            manpower = manpower.filter(start_date__lte=to_date)

        wb = Workbook()
        ws = wb.active
        ws.title = "Resources Report"
        bold_font = Font(bold=True)
        center_align = Alignment(horizontal="center")
        row = 1

        # Equipment
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
        ws.cell(row=row, column=1, value="Equipment").font = Font(bold=True, size=14)
        row += 1
        headers_eq = ["Project", "Name", "Category", "Quantity", "Condition", "Delivery Date"]
        for col_num, header in enumerate(headers_eq, 1):
            ws.cell(row=row, column=col_num, value=header).font = bold_font
            ws.cell(row=row, column=col_num).alignment = center_align
            ws.column_dimensions[get_column_letter(col_num)].width = max(len(header)+5, 15)
        row += 1
        for e in equipment:
            ws.append([e.project.project_name, e.name, e.category, e.quantity, e.condition, e.delivery_date.strftime("%Y-%m-%d")])
            row += 1
        row += 2

        # Manpower
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
        ws.cell(row=row, column=1, value="Manpower").font = Font(bold=True, size=14)
        row += 1
        headers_mp = ["Project", "Role", "Count", "Start Date"]
        for col_num, header in enumerate(headers_mp, 1):
            ws.cell(row=row, column=col_num, value=header).font = bold_font
            ws.cell(row=row, column=col_num).alignment = center_align
            ws.column_dimensions[get_column_letter(col_num)].width = max(len(header)+5, 15)
        row += 1
        for m in manpower:
            ws.append([m.project.project_name, m.role, m.count, m.start_date.strftime("%Y-%m-%d")])
            row += 1

        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        response = HttpResponse(
            buffer, content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        response["Content-Disposition"] = f'attachment; filename="resources_report_{datetime.date.today()}.xlsx"'
        return response

    except Exception as e:
        messages.error(request, f"Error exporting Excel: {e}")
        return redirect("reports:resources_report")


# ------------------ PDF ------------------
@login_required
@permission_required("reports.view_resourcesreport", raise_exception=True)
def resources_report_download_pdf(request):
    project_id = request.GET.get("project")
    from_date = request.GET.get("from_date")
    to_date = request.GET.get("to_date")

    if not project_id or project_id in ["None", ""]:
        messages.error(request, "Please select a project before downloading.")
        return redirect("reports:resources_report")

    try:
        # Filter resources
        equipment = filter_by_allowed_projects(Equipment.objects.filter(is_active=True), request.user)
        manpower = filter_by_allowed_projects(Manpower.objects.filter(is_active=True), request.user)

        if project_id:
            equipment = equipment.filter(project_id=int(project_id))
            manpower = manpower.filter(project_id=int(project_id))
        if from_date:
            equipment = equipment.filter(delivery_date__gte=from_date)
            manpower = manpower.filter(start_date__gte=from_date)
        if to_date:
            equipment = equipment.filter(delivery_date__lte=to_date)
            manpower = manpower.filter(start_date__lte=to_date)

        # -------------------------
        # Create PDF
        # -------------------------
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            leftMargin=36,
            rightMargin=36,
            topMargin=36,
            bottomMargin=36
        )

        styles = getSampleStyleSheet()
        elements = []

        title_style = ParagraphStyle(
            "Title", parent=styles["Heading1"], alignment=1, spaceAfter=12
        )
        section_style = ParagraphStyle(
            "Section", parent=styles["Heading2"], textColor=colors.HexColor("#0F5391"),
            spaceBefore=10, spaceAfter=5
        )
        cell_style = ParagraphStyle(
            "Cell", parent=styles["Normal"], fontSize=9, leading=12
        )

        # ---------- Title ----------
        elements.append(Paragraph("RESOURCES REPORT", title_style))

        # ---------- Equipment Table ----------
        elements.append(Paragraph("Equipment", section_style))
        eq_data = [["Project", "Name", "Category", "Quantity", "Condition", "Delivery Date"]]
        for e in equipment:
            eq_data.append([
                Paragraph(e.project.project_name, cell_style),
                Paragraph(e.name, cell_style),
                Paragraph(e.category, cell_style),
                Paragraph(str(e.quantity), cell_style),
                Paragraph(e.condition.capitalize(), cell_style),
                Paragraph(e.delivery_date.strftime("%Y-%m-%d"), cell_style)
            ])

        # Assign column widths proportional to full doc.width
        col_widths_eq = [
            doc.width * 0.2,  # Project
            doc.width * 0.25, # Name
            doc.width * 0.15, # Category
            doc.width * 0.1,  # Quantity
            doc.width * 0.15, # Condition
            doc.width * 0.15  # Delivery Date
        ]

        t_eq = Table(eq_data, colWidths=col_widths_eq, repeatRows=1)
        t_eq.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('ALIGN', (3,1), (3,-1), 'CENTER'),  # Quantity center
            ('ALIGN', (4,1), (4,-1), 'CENTER'),  # Condition center
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ]))
        elements.append(t_eq)
        elements.append(Spacer(1, 12))

        # ---------- Manpower Table ----------
        elements.append(Paragraph("Manpower", section_style))
        mp_data = [["Project", "Role", "Count", "Start Date"]]
        for m in manpower:
            mp_data.append([
                Paragraph(m.project.project_name, cell_style),
                Paragraph(m.role, cell_style),
                Paragraph(str(m.count), cell_style),
                Paragraph(m.start_date.strftime("%Y-%m-%d"), cell_style)
            ])

        col_widths_mp = [
            doc.width * 0.3,  # Project
            doc.width * 0.4,  # Role
            doc.width * 0.1,  # Count
            doc.width * 0.2   # Start Date
        ]

        t_mp = Table(mp_data, colWidths=col_widths_mp, repeatRows=1)
        t_mp.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('ALIGN', (2,1), (2,-1), 'CENTER'),  # Count center
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ]))
        elements.append(t_mp)

        # Build PDF
        doc.build(elements)
        buffer.seek(0)
        response = HttpResponse(buffer, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="resources_report_{datetime.date.today()}.pdf"'
        return response

    except Exception as e:
        messages.error(request, f"Error exporting PDF: {e}")
        return redirect("reports:resources_report")





# ------------------ HELPERS ------------------

def remove_paragraph_spacing(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()

    spacing = OxmlElement("w:spacing")
    spacing.set(ns.qn("w:before"), "0")
    spacing.set(ns.qn("w:after"), "0")
    spacing.set(ns.qn("w:line"), "240")
    spacing.set(ns.qn("w:lineRule"), "auto")

    pPr.append(spacing)


def style_header_cell(cell):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(10)


# ------------------ VIEW ------------------

@login_required
@permission_required("reports.view_resourcesreport", raise_exception=True)
def resources_report_download_word(request):

    project_id = request.GET.get("project")
    from_date = request.GET.get("from_date")
    to_date = request.GET.get("to_date")

    if not project_id:
        messages.error(request, "Please select a project before downloading.")
        return redirect("reports:resources_report")

    try:
        equipment = filter_by_allowed_projects(
            Equipment.objects.filter(is_active=True),
            request.user
        )
        manpower = filter_by_allowed_projects(
            Manpower.objects.filter(is_active=True),
            request.user
        )

        equipment = equipment.filter(project_id=project_id)
        manpower = manpower.filter(project_id=project_id)

        if from_date:
            equipment = equipment.filter(delivery_date__gte=from_date)
            manpower = manpower.filter(start_date__gte=from_date)

        if to_date:
            equipment = equipment.filter(delivery_date__lte=to_date)
            manpower = manpower.filter(start_date__lte=to_date)

        doc = Document()

        # ------------------ TITLE ------------------

        title = doc.add_heading("RESOURCES REPORT", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        remove_paragraph_spacing(title)

        # ------------------ EQUIPMENT ------------------

        eq_heading = doc.add_heading("Equipment", level=2)
        eq_heading.runs[0].bold = True
        remove_paragraph_spacing(eq_heading)

        table_eq = doc.add_table(rows=1, cols=6)
        table_eq.style = "Table Grid"
        table_eq.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers_eq = [
            "Project", "Name", "Category",
            "Quantity", "Condition", "Delivery Date"
        ]
        for i, h in enumerate(headers_eq):
            table_eq.rows[0].cells[i].text = h
            style_header_cell(table_eq.rows[0].cells[i])

        for e in equipment:
            row = table_eq.add_row().cells
            row[0].text = e.project.project_name
            row[1].text = e.name
            row[2].text = e.category
            row[3].text = str(e.quantity)
            row[4].text = e.condition.title()
            row[5].text = e.delivery_date.strftime("%Y-%m-%d")

            row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # collapse spacing after equipment table
        remove_paragraph_spacing(doc.paragraphs[-1])

        # ------------------ MANPOWER ------------------

        mp_heading = doc.add_heading("Manpower", level=2)
        mp_heading.runs[0].bold = True
        remove_paragraph_spacing(mp_heading)

        table_mp = doc.add_table(rows=1, cols=4)
        table_mp.style = "Table Grid"
        table_mp.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers_mp = ["Project", "Role", "Count", "Start Date"]
        for i, h in enumerate(headers_mp):
            table_mp.rows[0].cells[i].text = h
            style_header_cell(table_mp.rows[0].cells[i])

        for m in manpower:
            row = table_mp.add_row().cells
            row[0].text = m.project.project_name
            row[1].text = m.role
            row[2].text = str(m.count)
            row[3].text = m.start_date.strftime("%Y-%m-%d")

            row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        remove_paragraph_spacing(doc.paragraphs[-1])

        # ------------------ EXPORT ------------------

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response["Content-Disposition"] = (
            f'attachment; filename="resources_report_{datetime.date.today()}.docx"'
        )
        return response

    except Exception as e:
        messages.error(request, f"Error exporting Word document: {e}")
        return redirect("reports:resources_report")






# =============================
# FINANCE REPORTS
# =============================
@login_required
@permission_required("reports.view_financereport", raise_exception=True)
def finance_report(request):
    try:
        projects = get_allowed_projects(request.user)
        project_id = request.GET.get("project")
        from_date = request.GET.get("from_date")
        to_date = request.GET.get("to_date")

        is_filtered = bool(project_id or from_date or to_date)

        payments = PaymentCertificate.objects.none()
        transactions = FundTransaction.objects.none()

        if is_filtered:
            payments = filter_by_allowed_projects(
                PaymentCertificate.objects.filter(is_active=True),
                request.user,
                project_field="project"
            )
            transactions = filter_by_allowed_projects(
                FundTransaction.objects.filter(is_active=True),
                request.user,
                project_field="project"
            )

            if project_id:
                payments = payments.filter(project_id=project_id)
                transactions = transactions.filter(project_id=project_id)
            if from_date:
                payments = payments.filter(payment_date__gte=from_date)
                transactions = transactions.filter(date__gte=from_date)
            if to_date:
                payments = payments.filter(payment_date__lte=to_date)
                transactions = transactions.filter(date__lte=to_date)

        context = {
            "projects": projects,
            "payments": payments,
            "transactions": transactions,
            "filter_project": project_id,
            "filter_from": from_date,
            "filter_to": to_date,
            "is_filtered": is_filtered,
        }
        return render(request, "reports/finance/list.html", context)

    except Exception:
        logger.exception("Finance report error")
        messages.error(request, "Unable to load Finance Report.")
        return redirect("reports:finance_report")
    

# ------------------ EXCEL ------------------
@login_required
@permission_required("reports.view_financereport", raise_exception=True)
def finance_report_download_excel(request):
    project_id = request.GET.get("project")
    from_date = request.GET.get("from_date")
    to_date = request.GET.get("to_date")

    # Validate project filter
    if not project_id:
        messages.error(request, "Please select a project before downloading.")
        return redirect("reports:finance_report")

    payments = filter_by_allowed_projects(
        PaymentCertificate.objects.filter(is_active=True).order_by("-payment_date"),
        request.user
    )
    transactions = filter_by_allowed_projects(
        FundTransaction.objects.filter(is_active=True).order_by("date", "id"),
        request.user
    )

    if project_id:
        payments = payments.filter(project_id=int(project_id))
        transactions = transactions.filter(project_id=int(project_id))
    if from_date:
        payments = payments.filter(payment_date__gte=from_date)
        transactions = transactions.filter(date__gte=from_date)
    if to_date:
        payments = payments.filter(payment_date__lte=to_date)
        transactions = transactions.filter(date__lte=to_date)

    # Excel creation
    wb = Workbook()
    ws = wb.active
    ws.title = "Finance Report"
    row = 1
    bold_font = Font(bold=True)
    center_align = Alignment(horizontal="center")

    # Payment Certificates
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
    ws.cell(row=row, column=1, value="Payment Certificates").font = Font(bold=True, size=14)
    row += 1
    headers_pay = ["Project", "Certificate No", "Certified Amount", "Amount Paid", "Payment Date", "PV No"]
    for col_num, header in enumerate(headers_pay, 1):
        ws.cell(row=row, column=col_num, value=header).font = bold_font
        ws.cell(row=row, column=col_num).alignment = center_align
        ws.column_dimensions[get_column_letter(col_num)].width = max(len(header)+5, 15)
    row += 1
    for pay in payments:
        ws.cell(row=row, column=1, value=pay.project.project_name)
        ws.cell(row=row, column=2, value=pay.certificate_no)
        ws.cell(row=row, column=3, value=float(pay.certified_amount))
        ws.cell(row=row, column=4, value=float(pay.amount_paid))
        ws.cell(row=row, column=5, value=pay.payment_date.strftime("%Y-%m-%d"))
        ws.cell(row=row, column=6, value=pay.pv_no)
        row += 1
    row += 2

    # Fund Transactions
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
    ws.cell(row=row, column=1, value="Fund Utilization").font = Font(bold=True, size=14)
    row += 1
    headers_tx = ["Project", "Date", "Payee", "Type", "Amount Paid", "Balance After"]
    for col_num, header in enumerate(headers_tx, 1):
        ws.cell(row=row, column=col_num, value=header).font = bold_font
        ws.cell(row=row, column=col_num).alignment = center_align
        ws.column_dimensions[get_column_letter(col_num)].width = max(len(header)+5, 15)
    row += 1
    for tx in transactions:
        ws.cell(row=row, column=1, value=tx.project.project_name)
        ws.cell(row=row, column=2, value=tx.date.strftime("%Y-%m-%d"))
        ws.cell(row=row, column=3, value=tx.payee)
        ws.cell(row=row, column=4, value=tx.type)
        ws.cell(row=row, column=5, value=float(tx.amount_paid))
        ws.cell(row=row, column=6, value=float(tx.balance_after))
        row += 1

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    response = HttpResponse(
        buffer,
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    response["Content-Disposition"] = f'attachment; filename="finance_report_{datetime.date.today()}.xlsx"'
    return response


# ------------------ PDF ------------------
@login_required
@permission_required("reports.view_financereport", raise_exception=True)
def finance_report_download_pdf(request):
    project_id = request.GET.get("project")
    if not project_id:
        messages.error(request, "Please select a project before downloading.")
        return redirect("reports:finance_report")

    try:
        payments = filter_by_allowed_projects(PaymentCertificate.objects.filter(is_active=True), request.user)
        transactions = filter_by_allowed_projects(FundTransaction.objects.filter(is_active=True), request.user)

        if project_id:
            payments = payments.filter(project_id=int(project_id))
            transactions = transactions.filter(project_id=int(project_id))

        # -------------------------
        # Create PDF
        # -------------------------
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            leftMargin=36,
            rightMargin=36,
            topMargin=36,
            bottomMargin=18
        )

        styles = getSampleStyleSheet()
        elements = []

        title_style = ParagraphStyle(
            "Title", parent=styles["Heading1"], alignment=1, spaceAfter=12
        )
        section_style = ParagraphStyle(
            "Section", parent=styles["Heading2"], textColor=colors.HexColor("#0F5391"),
            spaceBefore=10, spaceAfter=5
        )
        cell_style = ParagraphStyle(
            "Cell", parent=styles["Normal"], fontSize=9, leading=12
        )

        # ---------- Title ----------
        elements.append(Paragraph("FINANCE REPORT", title_style))
        elements.append(Spacer(1, 12))

        # ---------- Payment Certificates ----------
        elements.append(Paragraph("Payment Certificates", section_style))
        data_payments = [["Project", "Certificate No", "Certified Amount", "Amount Paid", "Payment Date", "PV No"]]
        for pay in payments:
            data_payments.append([
                Paragraph(pay.project.project_name, cell_style),
                Paragraph(pay.certificate_no, cell_style),
                Paragraph(f"{pay.certified_amount:,.2f}", cell_style),
                Paragraph(f"{pay.amount_paid:,.2f}", cell_style),
                Paragraph(pay.payment_date.strftime("%Y-%m-%d"), cell_style),
                Paragraph(pay.pv_no, cell_style)
            ])

        col_widths_pay = [
            doc.width * 0.2,  # Project
            doc.width * 0.2,  # Certificate No
            doc.width * 0.15, # Certified Amount
            doc.width * 0.15, # Amount Paid
            doc.width * 0.15, # Payment Date
            doc.width * 0.15  # PV No
        ]

        t_pay = Table(data_payments, colWidths=col_widths_pay, repeatRows=1)
        t_pay.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('ALIGN', (2,1), (3,-1), 'RIGHT'),  # numeric columns
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ]))
        elements.append(t_pay)
        elements.append(Spacer(1, 12))

        # ---------- Fund Utilization ----------
        elements.append(Paragraph("Fund Utilization", section_style))
        data_tx = [["Project", "Date", "Payee", "Type", "Amount Paid", "Balance After"]]
        for tx in transactions:
            data_tx.append([
                Paragraph(tx.project.project_name, cell_style),
                Paragraph(tx.date.strftime("%Y-%m-%d"), cell_style),
                Paragraph(tx.payee, cell_style),
                Paragraph(tx.type, cell_style),
                Paragraph(f"{tx.amount_paid:,.2f}", cell_style),
                Paragraph(f"{tx.balance_after:,.2f}", cell_style),
            ])

        col_widths_tx = [
            doc.width * 0.2,  # Project
            doc.width * 0.15, # Date
            doc.width * 0.25, # Payee
            doc.width * 0.1,  # Type
            doc.width * 0.15, # Amount Paid
            doc.width * 0.15, # Balance After
        ]

        t_tx = Table(data_tx, colWidths=col_widths_tx, repeatRows=1)
        t_tx.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('ALIGN', (4,1), (5,-1), 'RIGHT'),  # numeric columns
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ]))
        elements.append(t_tx)

        # Build PDF
        doc.build(elements)
        buffer.seek(0)
        response = HttpResponse(buffer, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="finance_report_{datetime.date.today()}.pdf"'
        return response

    except Exception as e:
        messages.error(request, f"Error exporting PDF: {e}")
        return redirect("reports:finance_report")




# ------------------ HELPERS ------------------

def remove_paragraph_spacing(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()

    spacing = OxmlElement("w:spacing")
    spacing.set(ns.qn("w:before"), "0")
    spacing.set(ns.qn("w:after"), "0")
    spacing.set(ns.qn("w:line"), "240")
    spacing.set(ns.qn("w:lineRule"), "auto")

    pPr.append(spacing)


def style_header_cell(cell):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(10)


# ------------------ VIEW ------------------

@login_required
@permission_required("reports.view_financereport", raise_exception=True)
def finance_report_download_word(request):

    project_id = request.GET.get("project")
    if not project_id:
        messages.error(request, "Please select a project before downloading.")
        return redirect("reports:finance_report")

    payments = filter_by_allowed_projects(
        PaymentCertificate.objects.filter(is_active=True),
        request.user
    )
    transactions = filter_by_allowed_projects(
        FundTransaction.objects.filter(is_active=True),
        request.user
    )

    payments = payments.filter(project_id=project_id)
    transactions = transactions.filter(project_id=project_id)

    doc = Document()

    # ------------------ TITLE ------------------

    title = doc.add_heading("FINANCE REPORT", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    remove_paragraph_spacing(title)

    # ------------------ PAYMENT CERTIFICATES ------------------

    pay_heading = doc.add_heading("Payment Certificates", level=2)
    pay_heading.runs[0].bold = True
    remove_paragraph_spacing(pay_heading)

    table_pay = doc.add_table(rows=1, cols=6)
    table_pay.style = "Table Grid"
    table_pay.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers_pay = [
        "Project", "Certificate No",
        "Certified Amount", "Amount Paid",
        "Payment Date", "PV No"
    ]

    for i, h in enumerate(headers_pay):
        table_pay.rows[0].cells[i].text = h
        style_header_cell(table_pay.rows[0].cells[i])

    for pay in payments:
        row = table_pay.add_row().cells
        row[0].text = pay.project.project_name
        row[1].text = pay.certificate_no
        row[2].text = f"{pay.certified_amount:,.2f}"
        row[3].text = f"{pay.amount_paid:,.2f}"
        row[4].text = pay.payment_date.strftime("%Y-%m-%d")
        row[5].text = pay.pv_no

        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    remove_paragraph_spacing(doc.paragraphs[-1])

    # ------------------ FUND UTILIZATION ------------------

    tx_heading = doc.add_heading("Fund Utilization", level=2)
    tx_heading.runs[0].bold = True
    remove_paragraph_spacing(tx_heading)

    table_tx = doc.add_table(rows=1, cols=6)
    table_tx.style = "Table Grid"
    table_tx.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers_tx = [
        "Project", "Date", "Payee",
        "Type", "Amount Paid", "Balance After"
    ]

    for i, h in enumerate(headers_tx):
        table_tx.rows[0].cells[i].text = h
        style_header_cell(table_tx.rows[0].cells[i])

    for tx in transactions:
        row = table_tx.add_row().cells
        row[0].text = tx.project.project_name
        row[1].text = tx.date.strftime("%Y-%m-%d")
        row[2].text = tx.payee
        row[3].text = tx.type
        row[4].text = f"{tx.amount_paid:,.2f}"
        row[5].text = f"{tx.balance_after:,.2f}"

        row[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    remove_paragraph_spacing(doc.paragraphs[-1])

    # ------------------ EXPORT ------------------

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = (
        f'attachment; filename="finance_report_{datetime.date.today()}.docx"'
    )
    return response






# # =============================
# # QUALITY REPORTS
# # =============================
# @login_required
# @permission_required("reports.view_qualityreport", raise_exception=True)
# def quality_report(request):
#     try:
#         projects = get_allowed_projects(request.user)
#         project_id = request.GET.get("project")
#         from_date = request.GET.get("from_date")
#         to_date = request.GET.get("to_date")

#         is_filtered = bool(project_id or from_date or to_date)

#         material_tests = MaterialTest.objects.none()
#         work_approvals = WorkApproval.objects.none()

#         if is_filtered:
#             material_tests = filter_by_allowed_projects(
#                 MaterialTest.objects.filter(is_active=True),
#                 request.user,
#                 project_field="project"
#             )
#             work_approvals = filter_by_allowed_projects(
#                 WorkApproval.objects.filter(is_active=True),
#                 request.user,
#                 project_field="activity__project"
#             )

#             if project_id:
#                 material_tests = material_tests.filter(project_id=project_id)
#                 work_approvals = work_approvals.filter(activity__project_id=project_id)
#             if from_date:
#                 material_tests = material_tests.filter(test_date__gte=from_date)
#                 work_approvals = work_approvals.filter(approval_date__gte=from_date)
#             if to_date:
#                 material_tests = material_tests.filter(test_date__lte=to_date)
#                 work_approvals = work_approvals.filter(approval_date__lte=to_date)

#         context = {
#             "projects": projects,
#             "material_tests": material_tests,
#             "work_approvals": work_approvals,
#             "filter_project": project_id,
#             "filter_from": from_date,
#             "filter_to": to_date,
#             "is_filtered": is_filtered,
#         }
#         return render(request, "reports/quality/list.html", context)

#     except Exception:
#         logger.exception("Quality report error")
#         messages.error(request, "Unable to load Quality Report.")
#         return redirect("reports:quality_report")


# # -----------------------------
# # QUALITY REPORT EXCEL
# # -----------------------------
# @login_required
# @permission_required("reports.view_qualityreport", raise_exception=True)
# def quality_report_download_excel(request):
#     project_id = request.GET.get("project")
#     from_date = request.GET.get("from_date")
#     to_date = request.GET.get("to_date")

#     if not project_id:
#         messages.error(request, "Please select a project to download the report.")
#         return HttpResponse(status=400)

#     material_tests = filter_by_allowed_projects(MaterialTest.objects.filter(is_active=True), request.user, "project")
#     work_approvals = filter_by_allowed_projects(WorkApproval.objects.filter(is_active=True), request.user, "activity__project")

#     if project_id:
#         material_tests = material_tests.filter(project_id=project_id)
#         work_approvals = work_approvals.filter(activity__project_id=project_id)
#     if from_date:
#         material_tests = material_tests.filter(test_date__gte=from_date)
#         work_approvals = work_approvals.filter(approval_date__gte=from_date)
#     if to_date:
#         material_tests = material_tests.filter(test_date__lte=to_date)
#         work_approvals = work_approvals.filter(approval_date__lte=to_date)

#     response = HttpResponse(content_type="text/csv")
#     response["Content-Disposition"] = f'attachment; filename="quality_report_{project_id}.csv"'
#     writer = csv.writer(response)

#     # Material Tests
#     writer.writerow(["MATERIAL TESTS"])
#     writer.writerow(["Project", "Material", "Test Date", "Result", "Consultant"])
#     for t in material_tests:
#         writer.writerow([t.project.project_name, t.material_type, t.test_date, t.result, t.consultant])
#     writer.writerow([])

#     # Work Approvals
#     writer.writerow(["WORK APPROVALS"])
#     writer.writerow(["Project", "Activity", "Approved By", "Date", "Remarks"])
#     for a in work_approvals:
#         writer.writerow([
#             a.activity.project.project_name,
#             a.activity.name,
#             a.approved_by.username if a.approved_by else "",
#             a.approval_date,
#             a.remarks
#         ])

#     return response


# # -----------------------------
# # QUALITY REPORT PDF
# # -----------------------------
# @login_required
# @permission_required("reports.view_qualityreport", raise_exception=True)
# def quality_report_download_pdf(request):
#     project_id = request.GET.get("project")
#     from_date = request.GET.get("from_date")
#     to_date = request.GET.get("to_date")

#     if not project_id:
#         messages.error(request, "Please select a project to download the report.")
#         return redirect("reports:quality_report")

#     # -------------------------
#     # Filter data
#     # -------------------------
#     material_tests = filter_by_allowed_projects(MaterialTest.objects.filter(is_active=True), request.user, "project")
#     work_approvals = filter_by_allowed_projects(WorkApproval.objects.filter(is_active=True), request.user, "activity__project")

#     if project_id:
#         material_tests = material_tests.filter(project_id=project_id)
#         work_approvals = work_approvals.filter(activity__project_id=project_id)
#     if from_date:
#         material_tests = material_tests.filter(test_date__gte=from_date)
#         work_approvals = work_approvals.filter(approval_date__gte=from_date)
#     if to_date:
#         material_tests = material_tests.filter(test_date__lte=to_date)
#         work_approvals = work_approvals.filter(approval_date__lte=to_date)

#     # -------------------------
#     # Create PDF
#     # -------------------------
#     buffer = BytesIO()
#     doc = SimpleDocTemplate(
#         buffer,
#         pagesize=A4,
#         leftMargin=36,
#         rightMargin=36,
#         topMargin=36,
#         bottomMargin=18
#     )

#     styles = getSampleStyleSheet()
#     elements = []

#     title_style = ParagraphStyle(
#         "Title", parent=styles["Heading1"], alignment=1, spaceAfter=12
#     )
#     section_style = ParagraphStyle(
#         "Section", parent=styles["Heading2"], textColor=colors.HexColor("#0F5391"),
#         spaceBefore=10, spaceAfter=5
#     )
#     cell_style = ParagraphStyle(
#         "Cell", parent=styles["Normal"], fontSize=9, leading=12
#     )

#     # ---------- Title ----------
#     elements.append(Paragraph("QUALITY REPORT", title_style))
#     elements.append(Spacer(1, 12))

#     # ---------- Material Tests ----------
#     elements.append(Paragraph("Material Tests", section_style))
#     data_tests = [["Project", "Material Type", "Test Date", "Result", "Consultant"]]
#     for t in material_tests:
#         data_tests.append([
#             Paragraph(t.project.project_name, cell_style),
#             Paragraph(t.material_type, cell_style),
#             Paragraph(t.test_date.strftime("%Y-%m-%d"), cell_style),
#             Paragraph(t.result, cell_style),
#             Paragraph(t.consultant, cell_style)
#         ])

#     col_widths_tests = [
#         doc.width * 0.25,  # Project
#         doc.width * 0.15,  # Material Type
#         doc.width * 0.15,  # Test Date
#         doc.width * 0.10,  # Result
#         doc.width * 0.35,  # Consultant
#     ]

#     t_tests = Table(data_tests, colWidths=col_widths_tests, repeatRows=1)
#     t_tests.setStyle(TableStyle([
#         ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
#         ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
#         ('VALIGN', (0,0), (-1,-1), 'TOP'),
#         ('ALIGN', (2,1), (2,-1), 'RIGHT'),  # Test Date column right-aligned
#         ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
#     ]))
#     elements.append(t_tests)
#     elements.append(Spacer(1, 12))

#     # ---------- Work Approvals ----------
#     elements.append(Paragraph("Work Approvals", section_style))
#     data_approvals = [["Activity", "Approval Date", "Approved By", "Remarks"]]
#     for w in work_approvals:
#         data_approvals.append([
#             Paragraph(w.activity.name, cell_style),
#             Paragraph(w.approval_date.strftime("%Y-%m-%d"), cell_style),
#             Paragraph(w.approved_by.get_full_name() if w.approved_by else "-", cell_style),
#             Paragraph(w.remarks or "-", cell_style),
#         ])

#     col_widths_approvals = [
#         doc.width * 0.30,  # Activity
#         doc.width * 0.15,  # Approval Date
#         doc.width * 0.25,  # Approved By
#         doc.width * 0.30,  # Remarks
#     ]

#     t_approvals = Table(data_approvals, colWidths=col_widths_approvals, repeatRows=1)
#     t_approvals.setStyle(TableStyle([
#         ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
#         ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
#         ('VALIGN', (0,0), (-1,-1), 'TOP'),
#         ('ALIGN', (1,1), (1,-1), 'RIGHT'),  # Approval Date
#         ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
#     ]))
#     elements.append(t_approvals)

#     # -------------------------
#     # Build PDF
#     # -------------------------
#     doc.build(elements)
#     buffer.seek(0)
#     response = HttpResponse(buffer, content_type='application/pdf')
#     response['Content-Disposition'] = f'attachment; filename="quality_report_{datetime.date.today()}.pdf"'
#     return response


# # -----------------------------
# # QUALITY REPORT WORD
# # -----------------------------
# @login_required
# @permission_required("reports.view_qualityreport", raise_exception=True)
# def quality_report_download_word(request):
#     project_id = request.GET.get("project")
#     from_date = request.GET.get("from_date")
#     to_date = request.GET.get("to_date")

#     if not project_id:
#         messages.error(request, "Please select a project before downloading.")
#         return HttpResponse(status=400)

#     # Fetch filtered data
#     material_tests = filter_by_allowed_projects(
#         MaterialTest.objects.filter(is_active=True), request.user, "project"
#     )
#     work_approvals = filter_by_allowed_projects(
#         WorkApproval.objects.filter(is_active=True), request.user, "activity__project"
#     )

#     if project_id:
#         material_tests = material_tests.filter(project_id=project_id)
#         work_approvals = work_approvals.filter(activity__project_id=project_id)
#     if from_date:
#         material_tests = material_tests.filter(test_date__gte=from_date)
#         work_approvals = work_approvals.filter(approval_date__gte=from_date)
#     if to_date:
#         material_tests = material_tests.filter(test_date__lte=to_date)
#         work_approvals = work_approvals.filter(approval_date__lte=to_date)

#     # ------------------ CREATE DOCUMENT ------------------
#     doc = Document()

#     # ------------------ TITLE ------------------
#     title = doc.add_heading("QUALITY REPORT", level=1)
#     title.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     remove_paragraph_spacing(title)

#     # ------------------ MATERIAL TESTS ------------------
#     mt_heading = doc.add_heading("Material Tests", level=2)
#     mt_heading.runs[0].bold = True
#     remove_paragraph_spacing(mt_heading)

#     table_mt = doc.add_table(rows=1, cols=5)
#     table_mt.style = "Table Grid"
#     table_mt.alignment = WD_TABLE_ALIGNMENT.CENTER

#     headers_mt = ["Project", "Material", "Test Date", "Result", "Consultant"]
#     for i, h in enumerate(headers_mt):
#         table_mt.rows[0].cells[i].text = h
#         style_header_cell(table_mt.rows[0].cells[i])

#     for t in material_tests:
#         row = table_mt.add_row().cells
#         row[0].text = t.project.project_name
#         row[1].text = t.material_type
#         row[2].text = t.test_date.strftime("%Y-%m-%d")
#         row[3].text = t.result
#         row[4].text = t.consultant or "-"
#         # Right-align date column
#         row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

#     remove_paragraph_spacing(doc.paragraphs[-1])

#     # ------------------ WORK APPROVALS ------------------
#     wa_heading = doc.add_heading("Work Approvals", level=2)
#     wa_heading.runs[0].bold = True
#     remove_paragraph_spacing(wa_heading)

#     table_wa = doc.add_table(rows=1, cols=5)
#     table_wa.style = "Table Grid"
#     table_wa.alignment = WD_TABLE_ALIGNMENT.CENTER

#     headers_wa = ["Project", "Activity", "Approved By", "Approval Date", "Remarks"]
#     for i, h in enumerate(headers_wa):
#         table_wa.rows[0].cells[i].text = h
#         style_header_cell(table_wa.rows[0].cells[i])

#     for a in work_approvals:
#         row = table_wa.add_row().cells
#         row[0].text = a.activity.project.project_name
#         row[1].text = a.activity.name
#         row[2].text = a.approved_by.username if a.approved_by else "-"
#         row[3].text = a.approval_date.strftime("%Y-%m-%d")
#         row[4].text = a.remarks or "-"
#         # Right-align date column
#         row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

#     remove_paragraph_spacing(doc.paragraphs[-1])

#     # ------------------ EXPORT ------------------
#     buffer = BytesIO()
#     doc.save(buffer)
#     buffer.seek(0)

#     response = HttpResponse(
#         buffer,
#         content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#     )
#     response["Content-Disposition"] = (
#         f'attachment; filename="quality_report_{datetime.date.today()}.docx"'
#     )
#     return response



@login_required
@permission_required("reports.view_qualityreport", raise_exception=True)
def quality_report(request):
    try:
        projects = get_allowed_projects(request.user)
        project_id = request.GET.get("project")
        from_date = request.GET.get("from_date")
        to_date = request.GET.get("to_date")

        is_filtered = bool(project_id or from_date or to_date)

        material_tests = MaterialTest.objects.none()
        work_approvals = WorkApproval.objects.none()
        compliances = Compliance.objects.none()

        if is_filtered:
            material_tests = filter_by_allowed_projects(MaterialTest.objects.filter(is_active=True), request.user, "project")
            work_approvals = filter_by_allowed_projects(WorkApproval.objects.filter(is_active=True), request.user, "activity__project")
            compliances = filter_by_allowed_projects(Compliance.objects.filter(is_active=True), request.user, "project")

            if project_id:
                material_tests = material_tests.filter(project_id=project_id)
                work_approvals = work_approvals.filter(activity__project_id=project_id)
                compliances = compliances.filter(project_id=project_id)
            if from_date:
                material_tests = material_tests.filter(test_date__gte=from_date)
                work_approvals = work_approvals.filter(approval_date__gte=from_date)
                compliances = compliances.filter(expiry_date__gte=from_date)
            if to_date:
                material_tests = material_tests.filter(test_date__lte=to_date)
                work_approvals = work_approvals.filter(approval_date__lte=to_date)
                compliances = compliances.filter(expiry_date__lte=to_date)

        context = {
            "projects": projects,
            "material_tests": material_tests,
            "work_approvals": work_approvals,
            "compliances": compliances,
            "filter_project": project_id,
            "filter_from": from_date,
            "filter_to": to_date,
            "is_filtered": is_filtered,
        }
        return render(request, "reports/quality/list.html", context)

    except Exception:
        logger.exception("Quality report error")
        messages.error(request, "Unable to load Quality Report.")
        return redirect("reports:quality_report")


@login_required
@permission_required("reports.view_qualityreport", raise_exception=True)
def quality_report_download_excel(request):
    project_id = request.GET.get("project")
    from_date = request.GET.get("from_date")
    to_date = request.GET.get("to_date")

    if not project_id:
        return HttpResponse("Please select a project before downloading.", status=400)

    # Filter data
    material_tests = filter_by_allowed_projects(MaterialTest.objects.filter(is_active=True), request.user, "project")
    work_approvals = filter_by_allowed_projects(WorkApproval.objects.filter(is_active=True), request.user, "activity__project")
    compliances = filter_by_allowed_projects(Compliance.objects.filter(is_active=True), request.user, "project")

    if project_id:
        material_tests = material_tests.filter(project_id=project_id)
        work_approvals = work_approvals.filter(activity__project_id=project_id)
        compliances = compliances.filter(project_id=project_id)
    if from_date:
        material_tests = material_tests.filter(test_date__gte=from_date)
        work_approvals = work_approvals.filter(approval_date__gte=from_date)
        compliances = compliances.filter(expiry_date__gte=from_date)
    if to_date:
        material_tests = material_tests.filter(test_date__lte=to_date)
        work_approvals = work_approvals.filter(approval_date__lte=to_date)
        compliances = compliances.filter(expiry_date__lte=to_date)

    # Create CSV response
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="quality_report_{project_id}.csv"'
    writer = csv.writer(response)

    # Material Tests
    writer.writerow(["MATERIAL TESTS"])
    writer.writerow(["Project", "Material", "Test Date", "Result", "Consultant"])
    for t in material_tests:
        writer.writerow([t.project.project_name, t.material_type, t.test_date, t.result, t.consultant])
    writer.writerow([])

    # Work Approvals
    writer.writerow(["WORK APPROVALS"])
    writer.writerow(["Project", "Activity", "Approved By", "Approval Date", "Remarks"])
    for a in work_approvals:
        writer.writerow([
            a.activity.project.project_name,
            a.activity.name,
            a.approved_by.username if a.approved_by else "",
            a.approval_date,
            a.remarks or ""
        ])
    writer.writerow([])

    # Compliance
    writer.writerow(["COMPLIANCE"])
    writer.writerow(["Project", "Authority", "Registration No", "Status", "Expiry Date"])
    for c in compliances:
        writer.writerow([
            c.project.project_name,
            c.authority.name,
            c.registration_no,
            c.status,
            c.expiry_date
        ])

    return response


@login_required
@permission_required("reports.view_qualityreport", raise_exception=True)
def quality_report_download_pdf(request):
    project_id = request.GET.get("project")
    from_date = request.GET.get("from_date")
    to_date = request.GET.get("to_date")

    if not project_id:
        messages.error(request, "Please select a project to download the report.")
        return redirect("reports:quality_report")

    # -------------------------
    # Filter data
    # -------------------------
    material_tests = filter_by_allowed_projects(MaterialTest.objects.filter(is_active=True), request.user, "project")
    work_approvals = filter_by_allowed_projects(WorkApproval.objects.filter(is_active=True), request.user, "activity__project")
    compliances = filter_by_allowed_projects(Compliance.objects.filter(is_active=True), request.user, "project")

    if project_id:
        material_tests = material_tests.filter(project_id=project_id)
        work_approvals = work_approvals.filter(activity__project_id=project_id)
        compliances = compliances.filter(project_id=project_id)
    if from_date:
        material_tests = material_tests.filter(test_date__gte=from_date)
        work_approvals = work_approvals.filter(approval_date__gte=from_date)
        compliances = compliances.filter(expiry_date__gte=from_date)
    if to_date:
        material_tests = material_tests.filter(test_date__lte=to_date)
        work_approvals = work_approvals.filter(approval_date__lte=to_date)
        compliances = compliances.filter(expiry_date__lte=to_date)

    # -------------------------
    # Create PDF
    # -------------------------
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=18
    )

    styles = getSampleStyleSheet()
    elements = []

    title_style = ParagraphStyle(
        "Title", parent=styles["Heading1"], alignment=1, spaceAfter=12
    )
    section_style = ParagraphStyle(
        "Section", parent=styles["Heading2"], textColor=colors.HexColor("#0F5391"),
        spaceBefore=10, spaceAfter=5
    )
    cell_style = ParagraphStyle(
        "Cell", parent=styles["Normal"], fontSize=9, leading=12
    )

    # ---------- Title ----------
    elements.append(Paragraph("QUALITY REPORT", title_style))
    elements.append(Spacer(1, 12))

    # ---------- Material Tests ----------
    elements.append(Paragraph("Material Tests", section_style))
    data_tests = [["Project", "Material Type", "Test Date", "Result", "Consultant"]]
    for t in material_tests:
        data_tests.append([
            Paragraph(t.project.project_name, cell_style),
            Paragraph(t.material_type, cell_style),
            Paragraph(t.test_date.strftime("%Y-%m-%d"), cell_style),
            Paragraph(t.result, cell_style),
            Paragraph(t.consultant, cell_style)
        ])

    col_widths_tests = [
        doc.width * 0.25,  # Project
        doc.width * 0.15,  # Material Type
        doc.width * 0.15,  # Test Date
        doc.width * 0.10,  # Result
        doc.width * 0.35,  # Consultant
    ]

    t_tests = Table(data_tests, colWidths=col_widths_tests, repeatRows=1)
    t_tests.setStyle(TableStyle([
        ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
        ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('ALIGN', (2,1), (2,-1), 'RIGHT'),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
    ]))
    elements.append(t_tests)
    elements.append(Spacer(1, 12))

    # ---------- Work Approvals ----------
    elements.append(Paragraph("Work Approvals", section_style))
    data_approvals = [["Activity", "Approval Date", "Approved By", "Remarks"]]
    for w in work_approvals:
        data_approvals.append([
            Paragraph(w.activity.name, cell_style),
            Paragraph(w.approval_date.strftime("%Y-%m-%d"), cell_style),
            Paragraph(w.approved_by.get_full_name() if w.approved_by else "-", cell_style),
            Paragraph(w.remarks or "-", cell_style),
        ])

    col_widths_approvals = [
        doc.width * 0.30,  # Activity
        doc.width * 0.15,  # Approval Date
        doc.width * 0.25,  # Approved By
        doc.width * 0.30,  # Remarks
    ]

    t_approvals = Table(data_approvals, colWidths=col_widths_approvals, repeatRows=1)
    t_approvals.setStyle(TableStyle([
        ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
        ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('ALIGN', (1,1), (1,-1), 'RIGHT'),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
    ]))
    elements.append(t_approvals)
    elements.append(Spacer(1, 12))

    # ---------- Compliance ----------
    elements.append(Paragraph("Compliance", section_style))
    data_compliance = [["Project", "Authority", "Registration No", "Status", "Expiry Date"]]
    for c in compliances:
        data_compliance.append([
            Paragraph(c.project.project_name, cell_style),
            Paragraph(c.authority.name, cell_style),
            Paragraph(c.registration_no, cell_style),
            Paragraph(c.status, cell_style),
            Paragraph(c.expiry_date.strftime("%Y-%m-%d"), cell_style)
        ])

    col_widths_compliance = [
        doc.width * 0.25,  # Project
        doc.width * 0.25,  # Authority
        doc.width * 0.15,  # Registration No
        doc.width * 0.15,  # Status
        doc.width * 0.20,  # Expiry Date
    ]

    t_compliance = Table(data_compliance, colWidths=col_widths_compliance, repeatRows=1)
    t_compliance.setStyle(TableStyle([
        ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
        ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('ALIGN', (4,1), (4,-1), 'RIGHT'),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
    ]))
    elements.append(t_compliance)

    # -------------------------
    # Build PDF
    # -------------------------
    doc.build(elements)
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="quality_report_{datetime.date.today()}.pdf"'
    return response


@login_required
@permission_required("reports.view_qualityreport", raise_exception=True)
def quality_report_download_word(request):
    project_id = request.GET.get("project")
    from_date = request.GET.get("from_date")
    to_date = request.GET.get("to_date")

    if not project_id:
        messages.error(request, "Please select a project before downloading.")
        return HttpResponse(status=400)

    # -------------------------
    # Fetch filtered data
    # -------------------------
    material_tests = filter_by_allowed_projects(
        MaterialTest.objects.filter(is_active=True), request.user, "project"
    )
    work_approvals = filter_by_allowed_projects(
        WorkApproval.objects.filter(is_active=True), request.user, "activity__project"
    )
    compliances = filter_by_allowed_projects(
        Compliance.objects.filter(is_active=True), request.user, "project"
    )

    if project_id:
        material_tests = material_tests.filter(project_id=project_id)
        work_approvals = work_approvals.filter(activity__project_id=project_id)
        compliances = compliances.filter(project_id=project_id)
    if from_date:
        material_tests = material_tests.filter(test_date__gte=from_date)
        work_approvals = work_approvals.filter(approval_date__gte=from_date)
        compliances = compliances.filter(expiry_date__gte=from_date)
    if to_date:
        material_tests = material_tests.filter(test_date__lte=to_date)
        work_approvals = work_approvals.filter(approval_date__lte=to_date)
        compliances = compliances.filter(expiry_date__lte=to_date)

    # ------------------ CREATE DOCUMENT ------------------
    doc = Document()

    # ------------------ TITLE ------------------
    title = doc.add_heading("QUALITY REPORT", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    remove_paragraph_spacing(title)

    # ------------------ MATERIAL TESTS ------------------
    mt_heading = doc.add_heading("Material Tests", level=2)
    mt_heading.runs[0].bold = True
    remove_paragraph_spacing(mt_heading)

    table_mt = doc.add_table(rows=1, cols=5)
    table_mt.style = "Table Grid"
    table_mt.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers_mt = ["Project", "Material", "Test Date", "Result", "Consultant"]
    for i, h in enumerate(headers_mt):
        table_mt.rows[0].cells[i].text = h
        style_header_cell(table_mt.rows[0].cells[i])

    for t in material_tests:
        row = table_mt.add_row().cells
        row[0].text = t.project.project_name
        row[1].text = t.material_type
        row[2].text = t.test_date.strftime("%Y-%m-%d")
        row[3].text = t.result
        row[4].text = t.consultant or "-"
        # Right-align date column
        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    remove_paragraph_spacing(doc.paragraphs[-1])

    # ------------------ WORK APPROVALS ------------------
    wa_heading = doc.add_heading("Work Approvals", level=2)
    wa_heading.runs[0].bold = True
    remove_paragraph_spacing(wa_heading)

    table_wa = doc.add_table(rows=1, cols=5)
    table_wa.style = "Table Grid"
    table_wa.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers_wa = ["Project", "Activity", "Approved By", "Approval Date", "Remarks"]
    for i, h in enumerate(headers_wa):
        table_wa.rows[0].cells[i].text = h
        style_header_cell(table_wa.rows[0].cells[i])

    for a in work_approvals:
        row = table_wa.add_row().cells
        row[0].text = a.activity.project.project_name
        row[1].text = a.activity.name
        row[2].text = a.approved_by.username if a.approved_by else "-"
        row[3].text = a.approval_date.strftime("%Y-%m-%d")
        row[4].text = a.remarks or "-"
        row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    remove_paragraph_spacing(doc.paragraphs[-1])

    # ------------------ COMPLIANCE ------------------
    comp_heading = doc.add_heading("Compliance", level=2)
    comp_heading.runs[0].bold = True
    remove_paragraph_spacing(comp_heading)

    table_comp = doc.add_table(rows=1, cols=5)
    table_comp.style = "Table Grid"
    table_comp.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers_comp = ["Project", "Authority", "Registration No", "Status", "Expiry Date"]
    for i, h in enumerate(headers_comp):
        table_comp.rows[0].cells[i].text = h
        style_header_cell(table_comp.rows[0].cells[i])

    for c in compliances:
        row = table_comp.add_row().cells
        row[0].text = c.project.project_name
        row[1].text = c.authority.name
        row[2].text = c.registration_no
        row[3].text = c.status
        row[4].text = c.expiry_date.strftime("%Y-%m-%d")
        row[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    remove_paragraph_spacing(doc.paragraphs[-1])

    # ------------------ EXPORT ------------------
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = f'attachment; filename="quality_report_{datetime.date.today()}.docx"'
    return response




